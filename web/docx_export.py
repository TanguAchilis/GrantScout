"""
Word (.docx) export for GrantScout drafts.

Non-technical users (grant officers, founders) want a document they can open in
Word and finish — not Markdown. This builds a clean .docx from the drafted/
finalized sections, with the [ORG TO PROVIDE] markers visibly highlighted so the
"still needs a real figure" spots are impossible to miss.

Kept as a standalone, side-effect-free function (returns bytes) so it is unit-
testable without the web server (see tests/test_docx_export.py).
"""
from __future__ import annotations

import datetime
import io
import re

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt, RGBColor

MARKER = "[ORG TO PROVIDE]"
_TEAL = RGBColor(0x0F, 0x76, 0x6E)


def _strip_leading_heading(text: str) -> str:
    """Drop a leading markdown heading ('## ...') — the section title already
    carries it as a Word heading, so we avoid a duplicate."""
    return re.sub(r"^#{1,6}[^\n]*\n+", "", text or "")


def _add_body(doc: Document, body: str) -> None:
    """Add body text as paragraphs, splitting on blank lines, and highlight every
    [ORG TO PROVIDE] marker (bold + yellow) so it stands out for the reviewer."""
    for chunk in re.split(r"\n\s*\n", body):
        chunk = chunk.strip().replace("\n", " ")
        if not chunk:
            continue
        para = doc.add_paragraph()
        parts = chunk.split(MARKER)
        for i, part in enumerate(parts):
            if part:
                para.add_run(part)
            if i < len(parts) - 1:  # a marker sat between these parts
                run = para.add_run(MARKER)
                run.bold = True
                try:
                    run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                except Exception:  # highlight is cosmetic; never fail the export
                    pass


def build_docx(org_name: str, groups: list[dict], finalized: bool = False) -> bytes:
    """
    Build a .docx and return its bytes.

    `groups` is [{"grant_title": str, "sections": [{"title": str, "content": str}]}].
    Grouping by grant means each funder's application reads as its own titled block.
    """
    doc = Document()

    heading = doc.add_heading(level=0)
    hrun = heading.add_run(
        f"Grant Application Drafts{' (finalized)' if finalized else ''} — "
        f"{org_name or 'Your Organization'}"
    )
    hrun.font.color.rgb = _TEAL

    sub = doc.add_paragraph()
    srun = sub.add_run(
        f"Prepared by GrantScout · DRAFT · NOT SUBMITTED · {datetime.date.today().isoformat()}"
    )
    srun.italic = True
    srun.font.size = Pt(9)
    srun.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    disc = doc.add_paragraph()
    drun = disc.add_run(
        "GrantScout never submits applications. Review each draft, fill any "
        "[ORG TO PROVIDE] markers (highlighted below) with real figures, verify "
        "the funder's terms, and submit it yourself."
    )
    drun.italic = True

    for group in groups:
        doc.add_heading(group.get("grant_title") or "Grant", level=1)
        for section in group.get("sections", []):
            doc.add_heading(section.get("title") or "Section", level=2)
            _add_body(doc, _strip_leading_heading(section.get("content") or ""))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


__all__ = ["build_docx", "MARKER"]
